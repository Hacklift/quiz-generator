from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover - optional dependency at runtime
    fitz = None


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt"}
_WHITESPACE_PATTERN = re.compile(r"[ \t]+")
_MULTI_NEWLINE_PATTERN = re.compile(r"\n{3,}")


@dataclass
class ExtractedDocument:
    text: str
    source_document_name: str
    source_document_type: str
    title: str
    source_characters: int


def sanitize_document_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = _WHITESPACE_PATTERN.sub(" ", cleaned)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = _MULTI_NEWLINE_PATTERN.sub("\n\n", cleaned)
    return cleaned.strip()


def derive_document_title(filename: str | None, fallback_text: str = "") -> str:
    if filename:
        return Path(filename).stem.replace("_", " ").strip() or "Document Quiz"

    first_line = next(
        (line.strip() for line in fallback_text.splitlines() if line.strip()),
        "",
    )
    if first_line:
        return first_line[:80]
    return "Document Quiz"


def _extract_pdf_text(file_bytes: bytes) -> str:
    if fitz is not None:
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
            pages = [page.get_text("text") for page in pdf_document]
        return "\n".join(pages)

    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    return "\n".join(text for text in paragraphs if text)


def _extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode(errors="ignore")


def extract_text_from_bytes(
    *,
    file_bytes: bytes,
    filename: str,
) -> ExtractedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{extension}'. Supported types are PDF, DOCX, and TXT."
        )

    if extension == ".pdf":
        text = _extract_pdf_text(file_bytes)
        source_document_type = "pdf"
    elif extension == ".docx":
        text = _extract_docx_text(file_bytes)
        source_document_type = "docx"
    else:
        text = _extract_txt_text(file_bytes)
        source_document_type = "txt"

    cleaned_text = sanitize_document_text(text)
    if not cleaned_text:
        raise ValueError("No readable text was found in the uploaded document.")

    return ExtractedDocument(
        text=cleaned_text,
        source_document_name=os.path.basename(filename),
        source_document_type=source_document_type,
        title=derive_document_title(filename, cleaned_text),
        source_characters=len(cleaned_text),
    )


def extract_text_from_pasted_content(
    *,
    text: str,
    title: str | None = None,
) -> ExtractedDocument:
    cleaned_text = sanitize_document_text(text)
    if not cleaned_text:
        raise ValueError("Please provide text content to generate a quiz from.")

    document_title = title.strip() if title and title.strip() else derive_document_title(
        None,
        cleaned_text,
    )
    return ExtractedDocument(
        text=cleaned_text,
        source_document_name=document_title,
        source_document_type="text",
        title=document_title,
        source_characters=len(cleaned_text),
    )
